"""
Limpieza de documentos .docx: eliminar elementos duplicados y redundantes.

1. Separadores visuales (────────)
2. Etiquetas de lenguaje sueltas (PYTHON, JAVASCRIPT, TYPESCRIPT, etc.)
3. Párrafos "Código relevante:" sin contenido útil
4. Párrafos vacíos consecutivos (max 1)
5. BD/001: imágenes duplicadas del reemplazo masivo de placeholders
6. Tablas realmente duplicadas (mismo contenido completo)
7. Pies de foto duplicados consecutivos
8. Texto triplicado en headings (bug de python-docx: "TítuloTítuloTítulo")
"""

import sys, io, re
from pathlib import Path
from lxml import etree
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

DOCS = Path(r"C:\Users\freak\Preparacion Examenes\Documentos")
WP = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
DML = "http://schemas.openxmlformats.org/drawingml/2006/main"

# Language labels that appear as standalone paragraphs
LANG_LABELS = {
    'PYTHON', 'JAVASCRIPT', 'TYPESCRIPT', 'CSS', 'HTML', 'SQL', 'PHP',
    'JAVA', 'BASH', 'JSON', 'YAML', 'XML', 'RUBY', 'C', 'C++', 'RUST',
    'GO', 'KOTLIN', 'SWIFT', 'SCSS', 'SASS', 'LESS',
}

def has_image(elem):
    return bool(elem.findall(f'.//{{{DML}}}blip'))

def get_heading_info(elem):
    ppr = elem.find(f'{{{WP}}}pPr')
    if ppr is not None:
        pstyle = ppr.find(f'{{{WP}}}pStyle')
        if pstyle is not None:
            val = pstyle.get(f'{{{WP}}}val', '')
            if val.startswith('Heading'):
                try:
                    level = int(val.replace('Heading', ''))
                except ValueError:
                    level = 0
                return (level, ''.join(elem.itertext()).strip())
    return None

def get_full_table_text(tbl_elem):
    """Get full text content of a table for comparison."""
    return ''.join(tbl_elem.itertext()).strip()

def is_separator(txt):
    """Check if text is a separator line (────────)."""
    clean = txt.strip()
    if not clean:
        return False
    # All dashes, box-drawing chars, or similar
    return len(clean) >= 10 and all(c in '─━═—–-_' for c in clean)

def is_lang_label(txt):
    """Check if text is a standalone language label like 'PYTHON  PYTHON  PYTHON'."""
    clean = txt.strip()
    # Repeated language labels: "PYTHON  PYTHON  PYTHON" or just "PYTHON"
    words = [w for w in clean.split() if w]
    if not words:
        return False
    first = words[0].upper()
    if first in LANG_LABELS and all(w.upper() == first for w in words):
        return True
    return False

def fix_triple_heading_text(doc):
    """Fix headings where text appears 3x: 'TítuloTítuloTítulo' → 'Título'."""
    count = 0
    for p in doc.paragraphs:
        info = get_heading_info(p._element)
        if info:
            _, txt = info
            if len(txt) >= 6:
                third = len(txt) // 3
                if third >= 2 and txt[:third] == txt[third:2*third] == txt[2*third:3*third]:
                    # Triple text - fix it
                    correct = txt[:third]
                    # Clear all runs and set text
                    for run in p.runs:
                        run.text = ""
                    if p.runs:
                        p.runs[0].text = correct
                    else:
                        p.add_run(correct)
                    count += 1
    return count


print("=== Limpieza de documentos duplicados ===\n")

total_stats = {
    "separators": 0,
    "lang_labels": 0,
    "empty_consec": 0,
    "dup_tables": 0,
    "dup_images": 0,
    "dup_captions": 0,
    "triple_headings": 0,
    "codigo_relevante": 0,
}

for docx_path in sorted(DOCS.rglob("*.docx")):
    doc = Document(str(docx_path))
    body = doc.element.body
    rel = str(docx_path.relative_to(DOCS))

    to_remove = []
    stats = {k: 0 for k in total_stats}

    # ── Fix triple heading text ──
    stats["triple_headings"] = fix_triple_heading_text(doc)

    # ── Build element list ──
    elements = list(body)

    # ── Pass 1: Mark elements to remove ──
    prev_type = None  # 'empty', 'img', 'tbl', etc.
    prev_txt = None
    seen_tables = {}  # full_text → first element
    prev_was_empty = False
    prev_img_caption = None

    for i, child in enumerate(elements):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            txt = ''.join(child.itertext()).strip()
            is_img = has_image(child)
            heading = get_heading_info(child)

            if heading:
                prev_type = 'heading'
                prev_was_empty = False
                prev_img_caption = None
                continue

            if is_img:
                prev_type = 'img'
                prev_was_empty = False
                continue

            if not txt:
                if prev_was_empty:
                    to_remove.append(child)
                    stats["empty_consec"] += 1
                prev_was_empty = True
                prev_type = 'empty'
                continue

            prev_was_empty = False

            # Check separator
            if is_separator(txt):
                to_remove.append(child)
                stats["separators"] += 1
                continue

            # Check language label
            if is_lang_label(txt):
                to_remove.append(child)
                stats["lang_labels"] += 1
                continue

            # Check "Código relevante:" type paragraphs
            clean_lower = txt.lower().replace(' ', '')
            if clean_lower in ('códigorelevante:', 'codigorelevante:',
                               'códigorelevante:códigorelevante:códigorelevante:',
                               'codigorelevante:codigorelevante:codigorelevante:'):
                to_remove.append(child)
                stats["codigo_relevante"] += 1
                continue

            # Check duplicate caption (italic text that's identical to a previous one)
            # Detect patterns like "Catalogo de productos..." repeated
            if len(txt) > 15:
                # Check if this is a repeated caption (identical text nearby)
                pass  # handled by the BD/001 specific fix below

            prev_type = 'txt'
            prev_txt = txt

        elif tag == 'tbl':
            full_text = get_full_table_text(child)
            if full_text in seen_tables:
                to_remove.append(child)
                stats["dup_tables"] += 1
            else:
                seen_tables[full_text] = child
            prev_type = 'tbl'
            prev_was_empty = False

    # ── Pass 2 (BD/001 specific): Remove duplicate image+caption groups ──
    # In BD/001, the same 3 images (catalogo, producto, carrito) appear 8 times
    if "001-Proyecto tienda online" in rel:
        # Find all image paragraphs and their adjacent captions
        img_groups = []  # list of (img_elem, caption_elem_or_None, caption_text)
        elems = list(body)
        for i, child in enumerate(elems):
            if child in to_remove:
                continue
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'p' and has_image(child):
                # Check next sibling for caption
                caption_text = None
                caption_elem = None
                if i + 1 < len(elems):
                    nxt = elems[i + 1]
                    ntag = nxt.tag.split('}')[-1] if '}' in nxt.tag else nxt.tag
                    if ntag == 'p' and not has_image(nxt):
                        nxt_txt = ''.join(nxt.itertext()).strip()
                        if nxt_txt and len(nxt_txt) < 100:
                            caption_text = nxt_txt
                            caption_elem = nxt
                img_groups.append((child, caption_elem, caption_text))

        # Find duplicates: same caption appearing multiple times
        seen_captions = {}
        for img_elem, cap_elem, cap_text in img_groups:
            if cap_text:
                if cap_text in seen_captions:
                    # This is a duplicate image+caption pair
                    if img_elem not in to_remove:
                        to_remove.append(img_elem)
                        stats["dup_images"] += 1
                    if cap_elem and cap_elem not in to_remove:
                        to_remove.append(cap_elem)
                        stats["dup_captions"] += 1
                else:
                    seen_captions[cap_text] = True

    # ── Apply removals ──
    removed = 0
    for elem in to_remove:
        try:
            body.remove(elem)
            removed += 1
        except ValueError:
            pass

    if removed > 0:
        doc.save(str(docx_path))
        for k, v in stats.items():
            total_stats[k] += v
        details = [f"{v} {k}" for k, v in stats.items() if v > 0]
        print(f"  {rel}: eliminados {removed} ({', '.join(details)})")

print(f"\n=== Resumen ===")
for k, v in total_stats.items():
    if v > 0:
        print(f"  {k}: {v}")
print(f"  TOTAL eliminados: {sum(total_stats.values())}")
